import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING
from datetime import datetime
from together import Together
from llama_index.embeddings.together import TogetherEmbedding  # type: ignore
from llama_index.llms.together import TogetherLLM  # type: ignore

# Load API key from config
def load_api_key(config_path: str) -> str:
    if not os.path.exists(config_path):
        raise FileNotFoundError(f"The config file {config_path} does not exist.")
    with open(config_path, 'r') as file:
        config = json.load(file)
        
    return config.get("TOGETHER_API_KEY")

# Initialize Together client
def initialize_together_client(api_key: str) -> Together:
  
    return Together(api_key=api_key)

# Function to read text from a file
def read_text_file(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"No file found at {file_path}")
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()
    
# client = Together(api_key=load_api_key("config.json"))

# # Initialize embedding and generative models
# embed_model = TogetherEmbedding(model_name="togethercomputer/m2-bert-80M-8k-retrieval", api_key = load_api_key("config.json"))
# llm = TogetherLLM(model="meta-llama/Llama-3-8b-chat-hf", api_key=load_api_key("config.json"))

# # Read and process text from 'output.txt'
# file_text = read_text_file("output.txt")
# embeddings = embed_model.get_text_embedding(file_text)


def interpret_embeddings(embeddings):
    # Example: Pretend to analyze embeddings and return a summary
    return "return Team_Name, Key_Discussion_Points, Action_Items, Decisions_Made"

# # Construct a contextual prompt using interpreted embeddings
# interpreted_context = interpret_embeddings(embeddings)
# contextual_prompt = f"Based on the document's key themes, such as {interpreted_context}, provide a detailed analysis and summary."

# # Generate chat response using the LLM
# messages = [{"role": "user", "content": contextual_prompt + " " + file_text}]
# stream = client.chat.completions.create(
#     model="meta-llama/Llama-3-8b-chat-hf",
#     messages=messages,
#     stream=True,
# )


# # store the above response in a variable 
# response = ""
# for chunk in stream:
#     response += chunk.choices[0].delta.content or ""

# # print the response
# # print("response", response)

# Summarize text using the LLM
def summarize_text(text: str, client: Together) -> str:
    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        {
            "role": "user",
            "content": (
                f"Summarize the following meeting notes into a structured summary. Include the meeting lead(if mentioned), "
                "key discussion points, action items, and decisions made. Format the summary with clear headings: don't give me ** in output and also don't start with heading meeting summary as heading I will add it manually myself\n\n"
                f"{text}\n\n"
                "Meeting Lead: [Name](if mentioned or else skip this line)\n"
                "<h2>TeamName</h2>\n"
                "<h3>Key Discussion Points</h3>:\n"
                "-[Point 1]\n"
                "- [Point 2]\n"
                "<h3>Action Items</h3>:\n"
                "- [Item 1]\n"
                "- [Item 2]\n"
                "<h3>Decisions Made</h3>:\n"
                "- [Decision 1]\n"
                "- [Decision 2]\n"
            ),
        },
    ]

    response = client.chat.completions.create(
        model="meta-llama/Llama-3-8b-chat-hf",
        messages=messages,
        stream=False,
    )

    summary = response.choices[0].message.content.strip()
    return summary



def output_summary_to_docx(summary: str, file_path: str = "summary.docx"):
    doc = Document()

    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    today = datetime.today().strftime('%Y-%m-%d')

    # Heading with custom font and spacing settings
    heading_paragraph = doc.add_paragraph()
    run = heading_paragraph.add_run(f'Meeting Summary - {today}')
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading_paragraph.paragraph_format.space_after = Pt(12)
    heading_paragraph.paragraph_format.space_before = Pt(12)
    heading_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Process summary content
    lines = summary.split('\n')
    for line in lines:
        # Process and add paragraph for heading or normal text
        if '<h2>' in line or '<h3>' in line:
            line = line.replace('<h2>', '').replace('</h2>', '').replace('<h3>', '').replace('</h3>', '')
            style = 'Heading 2' if '<h2>' in line else 'Heading 3'
            paragraph = doc.add_paragraph(line, style=style)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        elif line.startswith('•'):
            # Properly handle pre-formatted bullet points
            paragraph = doc.add_paragraph(line)
            paragraph.style = 'List Bullet'
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        else:
            paragraph = doc.add_paragraph(line)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Save the document
    doc.save(file_path)
    print(f"Document saved as '{file_path}'")


# Main function
def main():
    config_path = "config.json"
    document_path = "output.txt"

    # Load API key
    api_key = load_api_key(config_path)

    # Initialize Together client
    client = initialize_together_client(api_key)

    # Read text from the file
    file_text = read_text_file(document_path)

    # Summarize text using the LLM
    summary = summarize_text(file_text, client)

    # Output summary to a .docx file
    output_summary_to_docx(summary)

    print(f"Summary: {summary}")

# Run the main function
if __name__ == "__main__":
    main()
